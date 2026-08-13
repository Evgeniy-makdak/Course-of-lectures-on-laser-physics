# -*- coding: utf-8 -*-
"""
Сборка печатной методички по 4 лекциям курса.
Порядок глав: бывшая Лекция 3 → 1; Лекция 1 → 2; Лекция 2 → 3; Лекция 4 → 4.
Текст = раскадровки без спикерских ремарок; со схемами, рисунками и таблицами.
"""
from __future__ import annotations

import ast
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, gray, black, white
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from _prepare_assets import prepare_all

ROOT = Path(r"C:\Users\Volkov\Desktop\Lecture-Unscheduled")
OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Методичка_Физика_лазеров_для_3D-технологий.docx"
PDF_PATH = OUT_DIR / "Методичка_Физика_лазеров_для_3D-технологий.pdf"
ASSETS = OUT_DIR / "assets"
# Версия для студентов (только уже прочитанные главы)
STUDENT_DIR = ROOT / "Методичка"
STUDENT_PDF_PATH = STUDENT_DIR / "Методичка_Физика_лазеров_для_3D-технологий.pdf"
AUTHOR = "Волков Е.В."

PREFACE = [
    (
        "Настоящая методичка собрана по материалам семинаров, которые автор читает для сотрудников АО "
        "«Лазерные Системы» в рамках курса по физике лазеров для практического применения в аддитивных "
        "(3D) технологиях в 2026 году."
    ),
    (
        "Текст глав соответствует содержанию лекций: сохранены определения, формулы, численные оценки, "
        "примеры и логика изложения."
    ),
    (
        "Методичка предназначена для самостоятельной работы студентов и инженеров аддитивных "
        "технологий как печатное сопровождение курса и будет дополняться после каждой новой прочитанной лекции."
    ),
]


# ── fonts / helpers ────────────────────────────────────────────────────────

def os_environ_windir() -> str:
    import os
    return os.environ.get("WINDIR", r"C:\Windows")


def _register_fonts():
    windir = Path(os_environ_windir())
    regular = windir / "Fonts" / "times.ttf"
    bold = windir / "Fonts" / "timesbd.ttf"
    italic = windir / "Fonts" / "timesi.ttf"
    if not regular.exists():
        regular = windir / "Fonts" / "arial.ttf"
        bold = windir / "Fonts" / "arialbd.ttf"
        italic = windir / "Fonts" / "ariali.ttf"
    pdfmetrics.registerFont(TTFont("Book", str(regular)))
    pdfmetrics.registerFont(TTFont("BookBold", str(bold if bold.exists() else regular)))
    if italic.exists():
        pdfmetrics.registerFont(TTFont("BookItalic", str(italic)))
        return "Book", "BookBold", "BookItalic"
    return "Book", "BookBold", "Book"


SPEAKER_LINE_RE = re.compile(
    r"^(?:ТАЙМИНГ\s*:|ДОСКА\s*/\s*ПАУЗА\s*:|□\s*Общее время|Общее время секции|"
    r"Длительность\s*:|Слайдов\s*:|Целевая аудитория\s*:|Файл презентации\s*:|"
    r"---\s*PAGE\s+\d+\s*---|---\s*КОНЕЦ|Страниц\s*:|Полная раскадровка|ПОЛНАЯ РАСКАДРОВКА)",
    re.I,
)
INLINE_TIMING_RE = re.compile(
    r"\s*[\|–-]?\s*\d{1,2}:\d{2}\s*[–-]\s*\d{1,2}:\d{2}(?:\s*\([^)]*\))?"
)
INLINE_HOLD_RE = re.compile(
    r"\s*\((?:держите на экране[^)]*|переключите[^)]*|≈\s*[^)]*минут[^)]*|"
    r"~\d+\s*мин[^)]*|быстро пройдите[^)]*)\)",
    re.I,
)
BOARD_CLOSING_RE = re.compile(
    r"^(?:Сейчас на доске|На доске (?:нарисуем|зафиксируем|запишем)|Запишем на доске|"
    r"Посмотрите на этот график)\b.*",
    re.I,
)

_SUP_MAP = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")
_SUB_MAP = str.maketrans("0123456789+-=()aehijklmnoprstuvx",
                         "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")


_SUP_LETTERS = str.maketrans({
    "n": "ⁿ", "N": "ⁿ", "i": "ⁱ", "x": "ˣ", "y": "ʸ", "z": "ᶻ",
    "a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ", "e": "ᵉ", "h": "ʰ",
    "j": "ʲ", "k": "ᵏ", "l": "ˡ", "m": "ᵐ", "o": "ᵒ", "p": "ᵖ",
    "r": "ʳ", "s": "ˢ", "t": "ᵗ", "u": "ᵘ", "v": "ᵛ", "w": "ʷ",
})


def _to_super(s: str) -> str:
    s = s.replace("−", "-").replace("–", "-")
    return s.translate(_SUP_MAP).translate(_SUP_LETTERS)


def _to_sub(s: str) -> str:
    s = s.replace("−", "-").replace("–", "-")
    return "".join(
        (ch.lower().translate(_SUB_MAP) if ch.lower() in "aehijklmnoprstuvx" else
         ch.translate(_SUB_MAP) if ch in "0123456789+-=()" else ch)
        for ch in s
    )


def typography_fix(text: str) -> str:
    """Normalize powers, indices, Greek nu, strip raw HTML sub/sup tags."""
    if not text:
        return text
    t = text.replace("\xa0", " ")

    # HTML tags -> unicode
    t = re.sub(r"<sup>\s*([^<]*?)\s*</sup>", lambda m: _to_super(m.group(1).replace("−", "-")), t, flags=re.I)
    t = re.sub(r"<sub>\s*([^<]*?)\s*</sub>", lambda m: _to_sub(m.group(1).replace("−", "-")), t, flags=re.I)
    # leftover angle-bracket artifacts
    t = t.replace("<sup>", "").replace("</sup>", "").replace("<sub>", "").replace("</sub>", "")

    # Summation formula (after tag conversion or raw HTML leftovers)
    t = re.sub(
        r"Σ(?:ₗ₌₀|ₗ=₀)?(?:ⁿ⁻¹)?\s*\(2l\+1\)",
        "Σₗ₌₀ⁿ⁻¹ (2l+1)",
        t,
    )

    # 10^-15 / 10^−15 / 10^{ -15 }
    t = re.sub(r"10\^\s*\{?\s*(−|-)?\s*(\d+)\s*\}?",
               lambda m: "10" + _to_super(("-" if m.group(1) else "") + m.group(2)), t)
    # letter^number (n^2, Z^2, v^2, r^2, including 2n^2)
    t = re.sub(r"([A-Za-z])\^\s*\{?\s*(−|-)?\s*(\d+)\s*\}?",
               lambda m: m.group(1) + _to_super(("-" if m.group(2) else "") + m.group(3)), t)
    t = re.sub(r"\b2n2\b", "2n²", t)
    t = re.sub(r"\b2N2\b", "2N²", t)

    # Energy/population indices: E1..E3, N1/N2, and underscore forms
    t = re.sub(r"\bE_([0-3])\b", lambda m: "E" + _to_sub(m.group(1)), t)
    t = re.sub(r"\bN_([12])\b", lambda m: "N" + _to_sub(m.group(1)), t)
    t = re.sub(r"\bE([0-3])\b", lambda m: "E" + _to_sub(m.group(1)), t)
    t = re.sub(r"\bN([12])\b", lambda m: "N" + _to_sub(m.group(1)), t)
    t = t.replace("m_l", "mₗ").replace("m_s", "mₛ").replace("N_max", "Nₘₐₓ")
    # Intensity / Bouguer forms
    t = t.replace("I_x = I_0 · e^(−α · x)", "I(x) = I₀ · exp(−α · x)")
    t = t.replace("I_x = I_0 · e^(-α · x)", "I(x) = I₀ · exp(−α · x)")
    t = re.sub(r"\bI_([0x])\b", lambda m: "I" + _to_sub(m.group(1)), t)
    t = re.sub(r"\bI0\b", "I₀", t)
    t = re.sub(
        r"\be\^\s*\(([^)]*)\)",
        lambda m: "exp(" + m.group(1).strip() + ")",
        t,
    )
    # leftover f-string artifacts
    t = re.sub(r"\{MOON_SPOT_KM:[^}]+\}", "3,8", t)
    t = re.sub(r"\{MOON_R_KM[^}]*\}", "384400", t)

    # Greek nu and multiplication
    t = t.replace("h*nu", "h·ν").replace("h·nu", "h·ν").replace("h ν", "h·ν")
    t = t.replace("h*ν", "h·ν")
    t = re.sub(r"(?<![A-Za-zА-Яа-я])nu(?![A-Za-zА-Яа-я])", "ν", t)

    # Delta E written as dE in formulas
    t = re.sub(r"(?<![A-Za-z])dE(?![A-Za-z])", "ΔE", t)
    t = t.replace("Delta E", "ΔE").replace("DeltaE", "ΔE")

    # normalize arrows / dots occasionally left as ASCII
    t = t.replace("<->", "↔").replace("->", "→")
    return t


def clean_heading(text: str) -> str:
    t = text.strip()
    t = re.sub(r"^СЛАЙД\s*:?\s*", "", t, flags=re.I)
    t = re.sub(r"^СЛАЙД\s+\d+\s*[-–—.:)]\s*", "", t, flags=re.I)
    t = INLINE_TIMING_RE.sub("", t)
    t = INLINE_HOLD_RE.sub("", t)
    # drop old numbering variants: "2.", "2 -", "2 —", "13 - ИТОГИ"
    t = re.sub(r"^\d+\s*[-–—.:)]\s*", "", t)
    t = re.sub(r"\s+", " ", t).strip(" |–-")
    return t


def clean_paragraph(text: str) -> str | None:
    t = text.strip()
    if not t or SPEAKER_LINE_RE.search(t) or re.fullmatch(r"\d{1,3}", t) or t.startswith("---"):
        return None
    if BOARD_CLOSING_RE.match(t):
        if t.lower().startswith("посмотрите на этот график"):
            t = re.sub(
                r"^Посмотрите на этот график\.\s*",
                "Рассмотрим график пороговой плотности энергии. ",
                t,
                flags=re.I,
            )
        else:
            return None
    t = INLINE_HOLD_RE.sub("", t)
    t = re.sub(r"\s+", " ", t).strip()
    reps = [
        (r"^Добрый день, коллеги\.\s*", ""),
        (r"^Добрый день\.\s*", ""),
        (r"^Поехали\.\s*", ""),
        (r"Проверьте сейчас:", "Проверим:"),
        (r"Давайте сделаем расчёт вместе", "Сделаем расчёт"),
        (r"Давайте вспомним", "Напомним"),
        (r"давайте вспомним", "напомним"),
        (r"Давайте посчитаем", "Посчитаем"),
        (r"давайте посчитаем", "посчитаем"),
        (r"Давайте посмотрим", "Рассмотрим"),
        (r"давайте посмотрим", "рассмотрим"),
        (r"Запишите гипотезу", "Сформулируйте гипотезу"),
        (r"Если есть вопросы по сегодняшнему материалу[^.]*\.\s*", ""),
        (r"Если вопросов нет,[^.]*\.\s*", ""),
        (r"короткий молчаливый повтор якорей на доске:[^.]*\.\s*", ""),
    ]
    for pat, rep in reps:
        t = re.sub(pat, rep, t)
    t = typography_fix(t.strip())
    return t or None


def load_script(py_path: Path):
    src = py_path.read_text(encoding="utf-8")
    tree = ast.parse(src)
    for node in tree.body:
        if isinstance(node, ast.Assign):
            for t in node.targets:
                if isinstance(t, ast.Name) and t.id == "SCRIPT":
                    return ast.literal_eval(node.value)
    raise RuntimeError(f"SCRIPT not found in {py_path}")


def paragraphs_from_script_items(script) -> list[tuple[str, list[str]]]:
    sections = []
    for item in script:
        title = clean_heading(item[0])
        paras = [cp for p in item[2] if (cp := clean_paragraph(p))]
        if paras:
            sections.append((title, paras))
    return sections


def normalize_sections(sections: list[tuple[str, list[str]]]) -> list[tuple[str, list[str]]]:
    """Unify first/last titles and numbering: '1. ВВЕДЕНИЕ' ... 'N. ПОДВЕДЕНИЕ ИТОГОВ'."""
    if not sections:
        return sections
    out = []
    for title, paras in sections:
        key = title.upper()
        if "ГОЛОСА ФИЗИКИ" in key:
            continue
        out.append((title, paras))
    if not out:
        return out

    # First section -> Введение
    out[0] = ("ВВЕДЕНИЕ", out[0][1])

    # Collect all summary-like sections into one finale
    summary_paras: list[str] = []
    body = []
    for title, paras in out[1:]:
        key = title.upper()
        if any(x in key for x in ("ПОДВЕДЕНИЕ ИТОГ", "ИТОГ", "ФИНАЛ", "ЗАКЛЮЧ")):
            summary_paras.extend(paras)
        else:
            body.append((title, paras))
    out = [("ВВЕДЕНИЕ", out[0][1])] + body
    if summary_paras:
        # de-dup adjacent identical paragraphs
        merged = []
        for p in summary_paras:
            if not merged or merged[-1] != p:
                merged.append(p)
        out.append(("ПОДВЕДЕНИЕ ИТОГОВ", merged))
    elif any(x in out[-1][0].upper() for x in ("ИТОГ", "ФИНАЛ", "ЗАКЛЮЧ")):
        out[-1] = ("ПОДВЕДЕНИЕ ИТОГОВ", out[-1][1])

    numbered = []
    for i, (title, paras) in enumerate(out, 1):
        core = re.sub(r"^\d+\s*[-–—.:)]\s*", "", title).strip().upper()
        core = typography_fix(core)
        numbered.append((f"{i}. {core}", paras))
    return numbered


# ── loaders ────────────────────────────────────────────────────────────────

def load_lecture3_sections() -> list[tuple[str, list[str]]]:
    script = load_script(ROOT / "Lecture-3-main" / "build_lecture_3.py")
    sections = paragraphs_from_script_items(script)

    # Drop quotes section
    sections = [(t, p) for t, p in sections if "ГОЛОСА" not in t.upper()]

    # Logical order for chapter 1
    order_keys = [
        "ТИТУЛ",
        "ПЛАН",
        "УСТРОЙСТВО АТОМА",
        "ПАУЛИ",
        "НАТРИ",
        "ПУТАНИЦ",
        "ДВЕ СХЕМ",
        "СОБСТВЕННОЕ СОСТОЯНИЕ",
        "ТОЧКИ НА ЛИНИ",
        "ВСЕГДА ЛИ",
        "УРОВНИ В РАЗНЫХ",
        "ПОДВЕДЕНИЕ",
        "ИТОГ",
    ]

    def rank(title: str) -> int:
        u = title.upper()
        for i, key in enumerate(order_keys):
            if key in u:
                return i
        return 50

    # Merge duplicate-ish titles carefully; keep all unique by rank then original
    sections = sorted(sections, key=lambda x: rank(x[0]))

    if sections:
        title0, paras0 = sections[0]
        new_paras = []
        for p in paras0:
            if "внеплановая лекция" in p.lower() or "не повторяет прошлое" in p.lower():
                continue
            if p.startswith("Сегодня мы сосредоточимся"):
                p = (
                    "Этой лекцией мы начинаем курс семинаров по физике лазеров "
                    "для практического применения в аддитивных (3D) технологиях. "
                    + p.replace("Сегодня мы сосредоточимся", "Мы сосредоточимся", 1)
                )
            new_paras.append(p)
        if not any("начинаем курс семинаров" in p for p in new_paras):
            new_paras.insert(
                0,
                "Этой лекцией мы начинаем курс семинаров по физике лазеров "
                "для практического применения в аддитивных (3D) технологиях.",
            )
        sections[0] = (title0, new_paras)

    adapted = []
    for title, paras in sections:
        ap = []
        for p in paras:
            p = p.replace("из прошлой лекции", "из базовой схемы лазерной физики")
            p = p.replace(
                "Вернёмся к базовой схеме из прошлой лекции",
                "Вернёмся к базовой схеме лазерной физики",
            )
            ap.append(p)
        adapted.append((title, ap))
    return normalize_sections(adapted)


def load_lecture1_sections() -> tuple[list[tuple[str, list[str]]], list[str]]:
    from docx import Document as Doc

    path = next((ROOT / "Lecture-1-main").glob("*раскадровка*.docx"))
    doc = Doc(path)
    raw_lines = [p.text.replace("\xa0", " ") for p in doc.paragraphs]

    glossary: list[str] = []
    sections: list[tuple[str, list[str]]] = []
    current_title = None
    current_paras: list[str] = []
    mode = "preamble"

    def flush():
        nonlocal current_title, current_paras
        if current_title and current_paras:
            sections.append((current_title, current_paras[:]))
        current_title = None
        current_paras = []

    for line in raw_lines:
        s = line.strip()
        if not s:
            continue
        if s.upper().startswith("ГЛОССАРИЙ"):
            mode = "glossary"
            continue
        if s.startswith("РАЗДЕЛ") or re.match(r"^СЛАЙД\s+\d+", s, re.I):
            mode = "body"
            flush()
            current_title = clean_heading(s)
            current_paras = []
            continue
        if mode == "glossary":
            if not SPEAKER_LINE_RE.search(s):
                glossary.append(re.sub(r"\s+", " ", s))
            continue
        if mode == "preamble" or SPEAKER_LINE_RE.search(s):
            continue
        cp = clean_paragraph(s)
        if cp:
            current_paras.append(cp)
            if cp.startswith("Закон Бугера") and "затухает с глубиной" in cp:
                current_paras.append("I(x) = I₀ · exp(−α · x)")
    flush()

    adapted = []
    for title, paras in sections:
        # Skip bare "РАЗДЕЛ ..." wrappers if they have little content? keep them if useful
        if title.upper().startswith("РАЗДЕЛ"):
            # convert section wrappers into ordinary titled blocks only if they have text
            pass
        ap = []
        for p in paras:
            p = p.replace(
                "Сегодня мы начинаем курс «Оптика и лазерная физика». Это не тот курс, где вас завалят формулами на первой же минуте. Мы пойдём другим путём: от понимания физической сути - к инженерным решениям.",
                "После знакомства со строением атома перейдём к вопросу о природе света — тоже без перегрузки формулами на первых шагах: от физической сути к инженерным решениям.",
            )
            p = p.replace("Именно об этом наша первая лекция.", "Именно об этом эта глава.")
            p = p.replace("Подведём итог нашей первой лекции.", "Подведём итог этой главы.")
            p = p.replace("которые я хочу, чтобы вы унесли с собой", "которые важно зафиксировать")
            p = p.replace(
                "На следующей лекции мы заглянем внутрь лазерного луча: разберём, что такое гауссов пучок, почему распределение интенсивности в пятне не равномерное, а колоколообразное; что такое поляризация и как она влияет на поглощение; и что такое когерентность и зачем она нужна - а когда она мешает.",
                "В следующей главе разберём, как из взаимодействия фотона с атомом возникает лазерный луч: вынужденное излучение, инверсия населённостей, резонатор и типы лазеров, применяемых в аддитивных технологиях.",
            )
            p = p.replace("будем обсуждать в следующей лекции", "будет важно в дальнейших главах")
            p = p.replace("ВЫВОДЫ ДЛЯ SLM", "ВЫВОДЫ ДЛЯ СЕЛЕКТИВНОГО ЛАЗЕРНОГО ПЛАВЛЕНИЯ")
            p = p.replace("для SLM", "для селективного лазерного плавления")
            p = p.replace("для вашего станка", "для установки селективного лазерного плавления")
            ap.append(p)
        # Drop pure section banners like "РАЗДЕЛ 1. ..." with almost no body? keep
        if title.upper().startswith("РАЗДЕЛ") and len(ap) == 0:
            continue
        if title.upper().startswith("РАЗДЕЛ"):
            continue  # section banners without body — skip; slides carry content
        adapted.append((title, ap))
    return normalize_sections(adapted), glossary


def _strip_html(text: str) -> str:
    t = text.replace("<b>", "").replace("</b>", "").replace("<i>", "").replace("</i>", "")
    return re.sub(r"<[^>]+>", "", t).strip()


def _eval_py_string(raw: str) -> str:
    raw = raw.strip()
    # Known constants from Lecture-2 builder (for leftover f-string placeholders)
    moon_r = 384_400
    moon_spot = 4.0 * moon_r / 400_000  # ≈ 3.8 km
    ns = {
        "MOON_R_KM": moon_r,
        "MOON_SPOT_KM": moon_spot,
        "MOON_SPOT_M": moon_spot * 1000,
    }
    if raw.startswith("f"):
        lit = raw[1:].strip()
        try:
            template = ast.literal_eval(lit)
        except Exception:
            template = lit.strip("'\"")
        # Evaluate simple {NAME:.1f} / {NAME} placeholders
        def _repl(m):
            expr = m.group(1)
            try:
                return format(eval(expr, {"__builtins__": {}}, ns))  # noqa: S307
            except Exception:
                if "MOON_SPOT_KM" in expr:
                    return f"{moon_spot:.1f}".replace(".", ",")
                if "MOON_R_KM" in expr:
                    return str(moon_r)
                return "3,8"
        return re.sub(r"\{([^{}]+)\}", _repl, template)
    try:
        return ast.literal_eval(raw)
    except Exception:
        parts = re.findall(r"'((?:\\'|[^'])*)'|\"((?:\\\"|[^\"])*)\"", raw)
        return "".join(a or b for a, b in parts)


def load_lecture2_sections() -> list[tuple[str, list[str]]]:
    src = (ROOT / "Lecture-2-main" / "build_lecture_2.py").read_text(encoding="utf-8")
    m = re.search(r"def build_pdf\(.*?:\n(.*?)(?=\ndef build_docx|\nif __name__)", src, re.S)
    body = m.group(1) if m else src
    sections: list[tuple[str, list[str]]] = []
    title = None
    paras: list[str] = []

    def flush():
        nonlocal title, paras
        if title and paras:
            sections.append((title, paras[:]))
        paras = []

    for m in re.finditer(
        r"Paragraph\(\s*'(РАЗДЕЛ[^']+|СЛАЙД[^']+)'\s*,|"
        r"slide_h3\(\s*'(СЛАЙД[^']+)'\s*\)|"
        r"para\(\s*(f?'(?:\\'|[^'])*')\s*\)|"
        r"formula\(\s*(f?'(?:\\'|[^'])*')\s*\)|"
        r"hint\(\s*((?:.|\n)*?)\s*\)\s*\)",
        body,
    ):
        if m.group(1) or m.group(2):
            flush()
            title = clean_heading(m.group(1) or m.group(2))
            continue
        raw = m.group(3) or m.group(4) or m.group(5)
        cp = clean_paragraph(_strip_html(_eval_py_string(raw)))
        if cp:
            paras.append(cp)
    flush()

    adapted = []
    for sec_title, sec_paras in sections:
        if sec_title.upper().startswith("РАЗДЕЛ"):
            continue
        ap = []
        for p in sec_paras:
            p = p.replace(
                "Добрый день, коллеги. Мы продолжаем наш цикл лекций по общей теме «Оптика и физика лазеров». Сегодня вторая лекция — «Физические основы работы лазера».",
                "После разбора строения атома и природы света перейдём к теме «Физические основы работы лазера».",
            )
            p = p.replace(
                "На прошлой лекции мы узнали, что свет — это поток фотонов.",
                "Ранее мы установили, что свет — это поток фотонов.",
            )
            p = p.replace("изученные на прошлой лекции", "введённые ранее")
            p = p.replace("Об этом — во второй части лекции.", "Об этом — далее в этой главе.")
            p = p.replace("во второй части лекции", "далее в этой главе")
            ap.append(p)
        if ap:
            adapted.append((sec_title, ap))
    return normalize_sections(adapted)


def load_lecture4_sections() -> list[tuple[str, list[str]]]:
    script = load_script(ROOT / "Lecture-4-main" / "build_lecture_4.py")
    sections = paragraphs_from_script_items(script)
    adapted = []
    for title, paras in sections:
        ap = []
        for p in paras:
            p = p.replace(
                "Мы продолжаем курс. На прошлых занятиях разобрали двойственную природу света и устройство лазера: за счёт чего излучение становится когерентным, направленным и узкополосным.",
                "В предыдущих главах мы разобрали строение атома, двойственную природу света и устройство лазера: за счёт чего излучение становится когерентным, направленным и узкополосным.",
            )
            p = p.replace("Тема лекции:", "Тема главы:")
            p = p.replace("План — пять связанных шагов.", "План главы — пять связанных шагов.")
            p = p.replace(
                "селективного лазерного плавления.",
                "селективного лазерного плавления (далее — СЛП).",
                1,
            )
            p = p.replace("вопрос на следующее занятие", "вопрос для самостоятельной проработки")
            p = p.replace(
                "Бонусный вопрос связывает сегодняшнюю тему с фокусировкой, которой займёмся дальше.",
                "Бонусный вопрос связывает тему этой главы с фокусировкой пучка — следующей логической темой курса.",
            )
            # Russianize residual English process jargon
            p = p.replace("lack of fusion", "несплавление")
            p = p.replace("keyhole-поры", "поры режима глубокого проплавления")
            p = p.replace("keyhole", "режим глубокого проплавления")
            p = p.replace("print-through", "пропечатывание насквозь")
            p = p.replace("шаг хэтча", "шаг штриховки")
            p = p.replace("green/blue", "зелёный/синий лазер")
            p = p.replace(
                "На следующем занятии разберём её вместе с геометрией пучка.",
                "Эту цепочку полезно разобрать далее вместе с геометрией пучка.",
            )
            p = p.replace(
                "На следующей лекции перейдём к геометрической оптике",
                "Далее по курсу логично перейти к геометрической оптике",
            )
            p = p.replace("якорь всей лекции", "якорь всей главы")
            p = p.replace("Центральный раздел лекции", "Центральный раздел главы")
            ap.append(p)
        adapted.append((title, ap))
    return normalize_sections(adapted)


# ── attach media ───────────────────────────────────────────────────────────

def _match(title: str, *keys: str) -> bool:
    u = title.upper()
    return any(k.upper() in u for k in keys)


def attach_media(chapters, catalog):
    """Attach figures/tables to sections; return chapters with rich section dicts."""
    fig_counters = {ch["number"]: 0 for ch in chapters}
    tab_counters = {ch["number"]: 0 for ch in chapters}

    def add_fig(ch_num, key):
        if key not in catalog["figures"]:
            return None
        fig_counters[ch_num] += 1
        meta = catalog["figures"][key]
        return {
            "type": "figure",
            "path": Path(meta["path"]),
            "caption": typography_fix(
                f"Рис. {ch_num}.{fig_counters[ch_num]}. {meta['caption']}"
            ),
        }

    def add_tab(ch_num, key):
        if key not in catalog["tables"]:
            return None
        tab_counters[ch_num] += 1
        meta = catalog["tables"][key]
        return {
            "type": "table",
            "caption": typography_fix(
                f"Таблица {ch_num}.{tab_counters[ch_num]}. {meta['caption']}"
            ),
            "headers": [typography_fix(h) for h in meta["headers"]],
            "rows": [[typography_fix(c) for c in row] for row in meta["rows"]],
        }

    # Mapping rules per chapter number
    rules = {
        1: [
            (("УСТРОЙСТВО АТОМА",), ["ch1_atom"]),
            (("ПАУЛИ",), ["ch1_pauli", "tab:ch1_quantum"]),
            (("НАТРИ",), ["ch1_sodium"]),
            (("ПУТАНИЦ",), ["ch1_two_schemes"]),
            (("ДВЕ СХЕМ", "СОБСТВЕНН"), ["ch1_two_schemes"]),
            (("ТОЧКИ",), ["ch1_points"]),
            (("ВСЕГДА ЛИ",), ["ch1_not_always"]),
        ],
        2: [
            (("ПОРОГОВОЙ ПЛОТНОСТИ", "ГРАФИК ПОРОГОВ"), ["ch2_threshold"]),
        ],
        3: [
            (("ЭНЕРГЕТИЧЕСКИЕ УРОВНИ",), ["ch3_spontaneous"]),
            (("ФОТОННАЯ ЛАВИНА", "MASER"), ["ch3_avalanche"]),
            (("ИНВЕРСНАЯ",), ["ch3_inversion_prob"]),
            (("РУБИНОВ",), ["ch3_ruby"]),
            (("РЕЗОНАТОР", "ОБРАТНАЯ СВЯЗЬ"), ["ch3_pump", "ch3_resonator"]),
            (("СРАВНЕНИЕ",), ["tab:ch3_laser_compare"]),
        ],
        4: [
            (("ЧАСТОТА", "ЭНЕРГИЯ ФОТОНА", "ДЛИНА ВОЛНЫ И ЭНЕРГИЯ"), ["ch4_energy"]),
            (("ШКАЛА",), ["ch4_spectrum"]),
            (("ДВА МЕХАНИЗМА", "ФОТОХИМИЧ"), ["ch4_mechanisms"]),
            (("ПОГЛОЩАТЕЛЬН", "СПЕКТРАЛЬН"), ["ch4_absorption", "tab:ch4_absorption"]),
            (("КРИТЕРИИ", "ВЫБОР ЛАЗЕРА"), ["ch4_choice"]),
        ],
    }

    rich_chapters = []
    for ch in chapters:
        rich_sections = []
        used_figs = set()
        for title, paras in ch["sections"]:
            blocks = [{"type": "text", "text": p} for p in paras]
            for keys, media_keys in rules.get(ch["number"], []):
                if _match(title, *keys):
                    for mk in media_keys:
                        if mk.startswith("tab:"):
                            item = add_tab(ch["number"], mk[4:])
                        else:
                            if mk in used_figs:
                                continue
                            item = add_fig(ch["number"], mk)
                            if item:
                                used_figs.add(mk)
                        if item:
                            insert_at = min(2, len(blocks))
                            blocks.insert(insert_at, item)
            rich_sections.append({"title": title, "blocks": blocks})
        # Ensure ch2 threshold appears if missed (attach to threshold section by fuzzy)
        if ch["number"] == 2 and "ch2_threshold" not in used_figs:
            for sec in rich_sections:
                if _match(sec["title"], "ПОРОГ", "ГРАФИК", "80 %"):
                    item = add_fig(2, "ch2_threshold")
                    if item:
                        sec["blocks"].insert(min(2, len(sec["blocks"])), item)
                        used_figs.add("ch2_threshold")
                    break
        rich_chapters.append({**ch, "sections": rich_sections})
    return rich_chapters


def build_book_model(catalog):
    ch1 = load_lecture3_sections()
    ch2, glossary = load_lecture1_sections()
    ch3 = load_lecture2_sections()
    ch4 = load_lecture4_sections()

    chapters = [
        {
            "number": 1,
            "title": "Строение атома и два языка лазерной физики",
            "subtitle": "Орбитали, квантовые числа и энергетические уровни E₁, E₂",
            "sections": ch1,
        },
        {
            "number": 2,
            "title": "Что такое свет, если забыть про формулы?",
            "subtitle": "Оптика и лазерная физика для инженеров аддитивных технологий",
            "sections": ch2,
        },
        {
            "number": 3,
            "title": "Физические основы работы лазера",
            "subtitle": "Вынужденное излучение, инверсия, резонатор и типы лазеров",
            "sections": ch3,
        },
        {
            "number": 4,
            "title": "Свет как поток фотонов: энергия, длина волны и выбор лазера",
            "subtitle": "Фотохимия и тепло, спектральное поглощение, критерии выбора источника",
            "sections": ch4,
        },
    ]
    rich = attach_media(chapters, catalog)
    return rich, list(PREFACE), glossary


# ── DOCX ───────────────────────────────────────────────────────────────────

def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if first_indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Cm(1.25)
    # Support **bold** fragments inside the paragraph
    parts = re.split(r"(\*\*.+?\*\*)", text)
    if len(parts) == 1 and not parts[0].startswith("**"):
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
        return p
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = p.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, italic=italic)
        else:
            run = p.add_run(part)
            set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_figure_docx(doc, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    # width ~14 cm
    run.add_picture(str(path), width=Cm(14.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    set_run_font(r, size=10, italic=True)


def add_table_docx(doc, headers, rows, caption: str):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(caption)
    set_run_font(r, size=10, italic=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def build_docx(chapters, preface, glossary, out_path: Path | None = None):
    out_path = out_path or DOCX_PATH
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, "ФИЗИКА ЛАЗЕРОВ ДЛЯ АДДИТИВНЫХ (3D) ТЕХНОЛОГИЙ",
             size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=16)
    add_para(doc, "Методическое пособие по материалам семинарских лекций",
             size=14, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=10)
    add_para(doc, f"Автор: {AUTHOR}",
             size=13, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=18)
    add_para(doc,
             "Для студентов и инженеров, осваивающих лазерные процессы в технологиях послойного синтеза",
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    doc.add_page_break()

    add_para(doc, "Предисловие", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
             first_indent=False, space_after=12)
    for p in preface:
        add_para(doc, p)
    doc.add_page_break()

    add_para(doc, "Содержание", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
             first_indent=False, space_after=12)
    for ch in chapters:
        add_para(doc, f"Глава {ch['number']}. {ch['title']}",
                 align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=2)
        add_para(doc, ch["subtitle"], size=11, italic=True,
                 align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=4)
        for sec in ch["sections"]:
            add_para(doc, f"    {sec['title']}", size=11,
                     align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=1)
        add_para(doc, "", space_after=6, first_indent=False)
    add_para(doc, "Приложение. Краткий глоссарий обозначений",
             align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    doc.add_page_break()

    for ch in chapters:
        add_para(doc, f"Глава {ch['number']}. {ch['title']}",
                 size=18, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=6)
        add_para(doc, ch["subtitle"], size=12, italic=True,
                 align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=14)
        for sec in ch["sections"]:
            add_para(doc, sec["title"], size=13, bold=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=8)
            for block in sec["blocks"]:
                if block["type"] == "text":
                    add_para(doc, block["text"])
                elif block["type"] == "figure":
                    if block["path"].exists():
                        add_figure_docx(doc, block["path"], block["caption"])
                elif block["type"] == "table":
                    add_table_docx(doc, block["headers"], block["rows"], block["caption"])
        doc.add_page_break()

    add_para(doc, "Приложение. Краткий глоссарий обозначений",
             size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=12)
    add_para(
        doc,
        "Глоссарий составлен по материалам главы о природе света и сохранён как справочник обозначений для всего курса.",
        italic=True, first_indent=False, space_after=10,
    )
    for g in glossary:
        if len(g) >= 3:
            add_para(doc, g, size=11, first_indent=False, space_after=4)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK DOCX: {out_path}")


# ── PDF ────────────────────────────────────────────────────────────────────

def pdf_escape(text: str, *, bold_font: str | None = None) -> str:
    # typography first (may introduce unicode), then escape XML specials;
    # **...** → bold (via explicit bold font name when provided)
    t = typography_fix(text)
    parts = re.split(r"(\*\*.+?\*\*)", t)
    out = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            inner = part[2:-2].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if bold_font:
                out.append(f'<font name="{bold_font}">{inner}</font>')
            else:
                out.append(f"<b>{inner}</b>")
        else:
            out.append(part.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    return "".join(out)


def build_pdf(chapters, preface, glossary, out_path: Path | None = None):
    out_path = out_path or PDF_PATH
    font, font_bold, font_italic = _register_fonts()

    def esc(text: str) -> str:
        return pdf_escape(text, bold_font=font_bold)

    styles = getSampleStyleSheet()
    cover = ParagraphStyle("Cover", parent=styles["Normal"], fontName=font_bold, fontSize=18,
                           leading=24, alignment=TA_CENTER, spaceAfter=16, textColor=HexColor("#1a1a2e"))
    cover_sub = ParagraphStyle("CoverSub", parent=styles["Normal"], fontName=font_italic, fontSize=12,
                               leading=16, alignment=TA_CENTER, spaceAfter=10, textColor=HexColor("#333333"))
    h1 = ParagraphStyle("H1", parent=styles["Normal"], fontName=font_bold, fontSize=14,
                        leading=18, alignment=TA_LEFT, spaceBefore=6, spaceAfter=8, textColor=HexColor("#1a1a2e"))
    h2 = ParagraphStyle("H2", parent=styles["Normal"], fontName=font_bold, fontSize=11.5,
                        leading=15, alignment=TA_LEFT, spaceBefore=10, spaceAfter=6, textColor=HexColor("#222222"))
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName=font, fontSize=10.5,
                          leading=14, alignment=TA_JUSTIFY, spaceAfter=6, firstLineIndent=14)
    body0 = ParagraphStyle("Body0", parent=body, firstLineIndent=0)
    note = ParagraphStyle("Note", parent=styles["Normal"], fontName=font_italic, fontSize=9.5,
                          leading=13, alignment=TA_CENTER, spaceAfter=10, textColor=HexColor("#444444"))
    toc = ParagraphStyle("TOC", parent=styles["Normal"], fontName=font, fontSize=11,
                         leading=15, alignment=TA_LEFT, spaceAfter=2)
    toc_sec = ParagraphStyle(
        "TOCSec", parent=toc, leftIndent=18, spaceAfter=1,
    )
    cell = ParagraphStyle("Cell", parent=styles["Normal"], fontName=font, fontSize=8,
                          leading=10, alignment=TA_LEFT)
    cell_h = ParagraphStyle("CellH", parent=cell, fontName=font_bold)

    story = []
    story.append(Spacer(1, 3 * cm))
    cover_author = ParagraphStyle(
        "CoverAuthor", parent=styles["Normal"], fontName=font, fontSize=12,
        leading=16, alignment=TA_CENTER, spaceAfter=14, textColor=HexColor("#222222"),
    )
    story.append(Paragraph(esc("ФИЗИКА ЛАЗЕРОВ ДЛЯ АДДИТИВНЫХ (3D) ТЕХНОЛОГИЙ"), cover))
    story.append(Paragraph(esc("Методическое пособие по материалам семинарских лекций"), cover_sub))
    story.append(Paragraph(esc(f"Автор: {AUTHOR}"), cover_author))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        esc("Для студентов и инженеров, осваивающих лазерные процессы в технологиях послойного синтеза"),
        cover_sub,
    ))
    story.append(PageBreak())

    story.append(Paragraph(esc("Предисловие"), h1))
    for p in preface:
        story.append(Paragraph(esc(p), body))
    story.append(PageBreak())

    story.append(Paragraph(esc("Содержание"), h1))
    for ch in chapters:
        story.append(Paragraph(esc(f"Глава {ch['number']}. {ch['title']}"), toc))
        story.append(Paragraph(esc(ch["subtitle"]), note))
        for sec in ch["sections"]:
            story.append(Paragraph(esc(sec["title"]), toc_sec))
    story.append(Paragraph(esc("Приложение. Краткий глоссарий обозначений"), toc))
    story.append(PageBreak())

    for ch in chapters:
        story.append(Paragraph(esc(f"Глава {ch['number']}. {ch['title']}"), h1))
        story.append(Paragraph(esc(ch["subtitle"]), note))
        for sec in ch["sections"]:
            story.append(Paragraph(esc(sec["title"]), h2))
            for block in sec["blocks"]:
                if block["type"] == "text":
                    story.append(Paragraph(esc(block["text"]), body))
                elif block["type"] == "figure" and block["path"].exists():
                    img = RLImage(str(block["path"]), width=14 * cm, height=9.5 * cm, kind="proportional")
                    story.append(Spacer(1, 6))
                    story.append(img)
                    story.append(Paragraph(esc(block["caption"]), note))
                elif block["type"] == "table":
                    story.append(Paragraph(esc(block["caption"]), note))
                    data = [[Paragraph(esc(h), cell_h) for h in block["headers"]]]
                    for row in block["rows"]:
                        data.append([Paragraph(esc(c), cell) for c in row])
                    col_w = 16.5 * cm / max(1, len(block["headers"]))
                    tbl = Table(data, colWidths=[col_w] * len(block["headers"]))
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#e8eef2")),
                        ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#666666")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 8))
        story.append(PageBreak())

    story.append(Paragraph(esc("Приложение. Краткий глоссарий обозначений"), h1))
    story.append(Paragraph(
        esc("Глоссарий составлен по материалам главы о природе света и сохранён как справочник обозначений для всего курса."),
        note,
    ))
    for g in glossary:
        if len(g) >= 3:
            story.append(Paragraph(esc(g), body0))

    def _page(canvas, doc_):
        canvas.saveState()
        canvas.setFont(font, 9)
        canvas.setFillColor(gray)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, str(canvas.getPageNumber()))
        canvas.restoreState()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=2.2 * cm, rightMargin=1.8 * cm,
        topMargin=1.6 * cm, bottomMargin=1.6 * cm,
        title="Физика лазеров для аддитивных (3D) технологий",
        author=AUTHOR,
    )
    doc.build(story, onFirstPage=_page, onLaterPages=_page)
    print(f"OK PDF: {out_path}")


def write_readme(chapters):
    lines = [
        "# Методичка: Физика лазеров для аддитивных (3D) технологий",
        "",
        "## Файлы",
        "",
        f"- `{DOCX_PATH.name}`",
        f"- `{PDF_PATH.name}`",
        "- `assets/` — рисунки и схемы",
        "",
        "## Главы",
        "",
    ]
    for ch in chapters:
        lines.append(f"{ch['number']}. {ch['title']}")
        for sec in ch["sections"]:
            n_fig = sum(1 for b in sec["blocks"] if b["type"] == "figure")
            n_tab = sum(1 for b in sec["blocks"] if b["type"] == "table")
            extra = []
            if n_fig:
                extra.append(f"{n_fig} рис.")
            if n_tab:
                extra.append(f"{n_tab} табл.")
            suffix = f" ({', '.join(extra)})" if extra else ""
            lines.append(f"   - {sec['title']}{suffix}")
    lines += ["", "Пересборка: `python build_print_handbook.py`", ""]
    (OUT_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")
    print("OK README")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    STUDENT_DIR.mkdir(parents=True, exist_ok=True)
    catalog = prepare_all()
    chapters, preface, glossary = build_book_model(catalog)
    n_fig = sum(1 for ch in chapters for sec in ch["sections"] for b in sec["blocks"] if b["type"] == "figure")
    n_tab = sum(1 for ch in chapters for sec in ch["sections"] for b in sec["blocks"] if b["type"] == "table")
    print(f"Chapters={len(chapters)}; figures={n_fig}; tables={n_tab}")
    for ch in chapters:
        print(f"  Ch{ch['number']}: {len(ch['sections'])} sections — first={ch['sections'][0]['title']}; last={ch['sections'][-1]['title']}")
    # Полная версия (все главы, включая ещё не прочитанные)
    build_docx(chapters, preface, glossary, DOCX_PATH)
    build_pdf(chapters, preface, glossary, PDF_PATH)
    # Версия для студентов: без глав, которые ещё не читались (сейчас — без гл. 4)
    student_chapters = [ch for ch in chapters if ch["number"] < 4]
    build_pdf(student_chapters, preface, glossary, STUDENT_PDF_PATH)
    write_readme(chapters)
    print("DONE")


if __name__ == "__main__":
    main()
